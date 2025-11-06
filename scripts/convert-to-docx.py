#!/usr/bin/env python3
"""
Convert the markdown course book to a formatted Word document (.docx)
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ python-docx library is required")
    print("   Please run: python3 -m pip install python-docx")
    sys.exit(1)

def add_heading(doc, text, level, style=None):
    """Add a heading with custom formatting"""
    if level == 1:
        para = doc.add_heading(text, level=1)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].font.size = Pt(24)
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    elif level == 2:
        para = doc.add_heading(text, level=2)
        para.runs[0].font.size = Pt(18)
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue
    elif level == 3:
        para = doc.add_heading(text, level=3)
        para.runs[0].font.size = Pt(14)
        para.runs[0].bold = True
    else:
        para = doc.add_heading(text, level=level)
    return para

def parse_markdown_to_docx(md_file_path, output_path):
    """Convert markdown to formatted Word document"""
    
    print(f"Reading {md_file_path}...")
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (will add spacing later)
        if not line:
            i += 1
            continue
        
        # Title (first # heading)
        if line.startswith('# ') and i < 5:
            title_text = line[2:].strip()
            add_heading(doc, title_text, level=1)
            
            # Add subtitle if next line is italic
            if i + 1 < len(lines) and lines[i+1].strip().startswith('*'):
                subtitle = lines[i+1].strip().strip('*').strip()
                para = doc.add_paragraph(subtitle)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].italic = True
                para.runs[0].font.size = Pt(14)
                i += 1
            
            # Add version if next line is bold
            if i + 1 < len(lines) and lines[i+1].strip().startswith('**'):
                version = lines[i+1].strip().strip('**').strip()
                para = doc.add_paragraph(version)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].font.size = Pt(11)
                i += 1
        
        # Part heading (# Part X:)
        elif line.startswith('# Part '):
            text = line[1:].strip()
            add_heading(doc, text, level=2)
        
        # Module heading (## Module X:)
        elif line.startswith('## Module '):
            text = line[2:].strip()
            add_heading(doc, text, level=2)
        
        # Lesson heading (### Lesson X.X:)
        elif line.startswith('### Lesson '):
            text = line[3:].strip()
            add_heading(doc, text, level=3)
        
        # Section heading (## for Table of Contents)
        elif line.startswith('## ') and 'Table of Contents' in line:
            text = line[2:].strip()
            add_heading(doc, text, level=2)
        
        # Regular heading (## or ###)
        elif line.startswith('## '):
            text = line[2:].strip()
            add_heading(doc, text, level=2)
        elif line.startswith('### '):
            text = line[3:].strip()
            add_heading(doc, text, level=3)
        elif line.startswith('#### '):
            text = line[4:].strip()
            add_heading(doc, text, level=4)
        
        # Horizontal rule
        elif line.startswith('---'):
            # Add some spacing
            doc.add_paragraph()
        
        # Bold text (**text**)
        elif line.startswith('**') and line.endswith('**'):
            text = line.strip('*').strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
        
        # Bullet list item (- or *)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove markdown formatting from list items
            text = text.replace('**', '').replace('`', '')
            para = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered list item
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = text.replace('**', '').replace('`', '')
            para = doc.add_paragraph(text, style='List Number')
        
        # Table
        elif line.startswith('|') and '---' not in line:
            # Start of a table
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if '---' not in row_line:
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    table_data.append(cells)
                i += 1
            
            if table_data:
                # Create table
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        if row_idx < len(table_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                i -= 1  # Adjust because we already advanced
        
        # Regular paragraph
        else:
            # Remove markdown formatting
            text = line
            text = text.replace('**', '').replace('*', '').replace('`', '')
            
            # Handle italic sections
            if text.startswith('*') and text.endswith('*'):
                para = doc.add_paragraph(text.strip('*'))
                para.runs[0].italic = True
            else:
                # Check if it's a label like "**Key Concept**"
                if text.startswith('**') and text.endswith('**'):
                    para = doc.add_paragraph()
                    run = para.add_run(text.strip('*'))
                    run.bold = True
                    run.font.size = Pt(12)
                elif text.startswith('**') and ':**' in text:
                    # Bold label like "**Overview:** text"
                    parts = text.split(':**', 1)
                    para = doc.add_paragraph()
                    run1 = para.add_run(parts[0].strip('*') + ':')
                    run1.bold = True
                    if len(parts) > 1:
                        run2 = para.add_run(' ' + parts[1].strip())
                        run2.font.size = Pt(11)
                elif text.startswith('💡'):
                    para = doc.add_paragraph(text)
                    para.runs[0].font.size = Pt(11)
                else:
                    # Regular paragraph
                    para = doc.add_paragraph(text)
                    para.runs[0].font.size = Pt(11)
        
        i += 1
    
    # Save document
    print(f"Saving to {output_path}...")
    doc.save(output_path)
    print(f"✅ Successfully created {output_path}")

if __name__ == '__main__':
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    md_file = project_root / 'COMPLETE-COURSE-BOOK.md'
    
    # Desktop path
    desktop = Path.home() / 'Desktop'
    output_file = desktop / 'Power-BI-Mastery-Course-Book.docx'
    
    if not md_file.exists():
        print(f"❌ File not found: {md_file}")
        sys.exit(1)
    
    parse_markdown_to_docx(md_file, output_file)

